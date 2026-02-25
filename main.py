import os
import shutil
import logging
import asyncio
from datetime import datetime
from dotenv import load_dotenv

# FastAPI Imports
from fastapi import FastAPI, UploadFile, File, Form, BackgroundTasks, HTTPException
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware

# Your existing logic imports
import PyPDF2
from google import genai
from playwright.async_api import async_playwright
from docx import Document
from docx.shared import Pt, Inches, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

# ==========================================
# CONFIGURATION & LOGGING
# ==========================================
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[logging.StreamHandler()]
)

load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
# Check if we are in DEV_MODE (defaults to False if not found)
DEV_MODE = os.getenv("DEV_MODE", "false").lower() == "true"

if not GEMINI_API_KEY and not DEV_MODE:
    logging.critical("GEMINI_API_KEY not found in .env file.")
    exit(1)

if not DEV_MODE:
    client = genai.Client(api_key=GEMINI_API_KEY)
else:
    logging.info("🛠️ Running in DEV MODE. AI and Web Scraping will be bypassed.")

# Setup App & CORS
app = FastAPI(title="Resume Tailor API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], 
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ==========================================
# CORE LOGIC FUNCTIONS
# ==========================================

def extract_text_from_pdf(pdf_path):
    logging.info(f"Starting text extraction from PDF: {pdf_path}")
    text = ""
    try:
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        logging.info(f"Successfully extracted {len(text)} characters from PDF.")
        return text
    except Exception as e:
        logging.error(f"Error reading PDF {pdf_path}: {e}")
        return None

async def fetch_jd_from_url(url):
    logging.info(f"Launching headless browser to fetch JD from: {url}")
    try:
        async with async_playwright() as p:
            browser = await p.chromium.launch(headless=True)
            page = await browser.new_page()
            await page.goto(url, wait_until="networkidle", timeout=30000) 
            text = await page.locator("body").inner_text()
            await browser.close()
            logging.info(f"Successfully fetched {len(text)} characters from JD URL.")
            return text
    except Exception as e:
        logging.error(f"Error fetching JD from {url}: {e}")
        return None

def tailor_resume_with_ai(resume_text, jd_text):
    logging.info("Sending prompt to Gemini API...")
    prompt = f"""
    You are an expert technical recruiter. I am applying for a Software Development Engineer (SDE-I) role. 
    Rewrite my resume to highly align with the JD, keeping it factual and strictly based on my actual experience.
    
    CRITICAL OUTPUT RULES:
    1. DO NOT output any conversational text, pleasantries, or preambles. Start the response immediately with Line 1.
    
    CRITICAL CONTENT & ORDER RULES:
    1. You MUST structure the resume sections in this EXACT order: 
       PROFESSIONAL SUMMARY, EXPERIENCE, PROJECTS, TECHNICAL SKILLS, CERTIFICATIONS, ACHIEVEMENTS, EDUCATION.
       If any of these sections are missing from my original resume, omit them completely. Do not invent data.
    2. DO NOT delete any Projects, Education, Experience, or Certifications. Include all technical sections.
    3. DO NOT include a "Languages" section for spoken/natural languages (e.g., English, Tamil, Hindi).
    4. Compress descriptions to fit on a single A4 page. Max 3 bullet points per job/project.
    5. Remove LeetCode from the contact information.
    6. Explicitly include the exact title "Software Engineer - Backend" or "Backend Developer" in the Professional Summary.
    7. Maximize keyword density naturally for the following ATS keywords throughout the summary and experience bullet points: 'API development', 'microservices', 'authentication', 'security', 'OAuth2', 'JWT', 'Spring Boot', 'Node.js', 'PostgreSQL', 'Redis', 'Docker', 'Hibernate', and 'Keycloak'.
    8. Consistently quantify achievements in bullet points with metrics (e.g., 'reduced latency by 70%', 'supported 1M+ users', '~2s response time').
    9. SKILLS SECTION: Ensure all the above keywords are explicitly listed in the skills section if applicable, but keep it readable (max 5-8 per category).
    10. EDUCATION SECTION: ONLY include the institution name, date, and degree. DO NOT add any foundational coursework, descriptions, or bullet points.
    
    STRICT FORMATTING & BOLDING RULES (CRITICAL):
    1. Line 1: My Name
    2. Line 2 (and Line 3 if necessary): Contact info separated by pipes using explicit labels (e.g., Email: <email> | Phone: <phone> | LinkedIn: <link> | GitHub: <link> | Portfolio: <link>). ONLY include details that exist in the original resume. Omit missing items completely.
    3. CRITICAL SPACE-SAVING RULE: For all URLs in the contact info, you MUST strip "https://" and "www." to save space (e.g., output "linkedin.com/in/dhanush-karthik").
    4. Section Headers MUST start with "SECTION: " (e.g., "SECTION: EXPERIENCE", "SECTION: CERTIFICATIONS").
    5. For Experience and Education headers, use a double-pipe "||" to split the title/company from the date. (Example: Queen Mary University of London || 2024 – 2025).
    6. For Project titles and Award titles, wrap the whole line in ** to make it bold.
    7. In the Technical Skills section, ONLY wrap the category name in **. Leave the actual skills plain text.
    8. Inside Experience and Projects bullet points, use **bold** text to highlight key technologies, metrics, and ATS keywords.
    9. Start bullet points with a single dash "- ". This applies to Certifications as well; list them as single-line bullet points.

    CURRENT RESUME:
    {resume_text}

    JOB DESCRIPTION:
    {jd_text}
    """
    try:
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=prompt
        )
        logging.info("Successfully received tailored resume content from Gemini.")
        return response.text
    except Exception as e:
        logging.error(f"Error generating tailored resume via API: {e}")
        return None

def add_bottom_border(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnvPPr', 'w:info', 'w:wpPr'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)

def set_compact_spacing(paragraph, space_after=2):
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.space_before = Pt(0)

def add_markdown_run(paragraph, text, base_font_name='Calibri', base_font_size=10):
    parts = text.split('**')
    for i, part in enumerate(parts):
        if not part: continue
        run = paragraph.add_run(part)
        run.font.name = base_font_name
        run.font.size = Pt(base_font_size)
        if i % 2 != 0: run.bold = True

def create_styled_docx(text, output_filename):
    logging.info("Formatting data into DOCX...")
    try:
        doc = Document()
        section = doc.sections[0]
        section.page_width, section.page_height = Mm(210), Mm(297)
        section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(0.5)
        right_tab_stop = Inches(7.27)

        lines = [line.strip() for line in text.split('\n') if line.strip()]
        if not lines: 
            logging.warning("No text lines found to format into DOCX.")
            return

        # Name
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_compact_spacing(name_para, space_after=2)
        name_run = name_para.add_run(lines[0].replace('**', ''))
        name_run.bold = True
        name_run.font.name, name_run.font.size = 'Calibri', Pt(16)

        # Contact
        content_start_idx = 1
        for i in range(1, len(lines)):
            if lines[i].startswith("SECTION:"):
                content_start_idx = i
                break
                
        for i in range(1, content_start_idx):
            contact_para = doc.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_compact_spacing(contact_para, space_after=2)
            contact_run = contact_para.add_run(lines[i].replace('**', ''))
            contact_run.font.name, contact_run.font.size = 'Calibri', Pt(10)
            if i == content_start_idx - 1:
                contact_para.paragraph_format.space_after = Pt(8)

        current_section = ""
        cert_table, cert_cells = None, []

        for line in lines[content_start_idx:]:
            if not line.strip(): continue

            if line.startswith("SECTION:"):
                header_text = line.replace("SECTION:", "").strip().upper()
                current_section = header_text
                p = doc.add_paragraph()
                set_compact_spacing(p, space_after=2)
                p.paragraph_format.space_before = Pt(6)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER  
                run = p.add_run(header_text)
                run.bold, run.font.name, run.font.size = True, 'Calibri', Pt(11)
                add_bottom_border(p)
                
                if "CERTIFICATION" in current_section:
                    cert_table = doc.add_table(rows=0, cols=2)
                    cert_table.autofit, cert_cells = False, []
                continue

            if "CERTIFICATION" in current_section:
                clean_line = line.lstrip('-•* ')
                if not cert_cells:
                    row = cert_table.add_row()
                    row.cells[0].width, row.cells[1].width = Inches(3.635), Inches(3.635)
                    cert_cells = [row.cells[0], row.cells[1]]
                cell = cert_cells.pop(0)
                p = cell.paragraphs[0]
                set_compact_spacing(p, space_after=2)
                add_markdown_run(p, "• " + clean_line)
                continue
            
            elif line.startswith('-') or line.startswith('•'):
                clean_line = line.lstrip('-•* ')
                p = doc.add_paragraph(style='List Bullet')
                set_compact_spacing(p, space_after=2)
                p.paragraph_format.left_indent = Inches(0.25)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  
                add_markdown_run(p, clean_line)
            
            elif '||' in line:
                left_part, right_part = [part.strip() for part in line.split('||', 1)]
                p = doc.add_paragraph()
                set_compact_spacing(p, space_after=2)
                p.paragraph_format.tab_stops.add_tab_stop(right_tab_stop, WD_TAB_ALIGNMENT.RIGHT)
                run_left = p.add_run(left_part.replace('**', '') + '\t')
                run_left.bold, run_left.font.name, run_left.font.size = True, 'Calibri', Pt(10.5)
                run_right = p.add_run(right_part.replace('**', ''))
                run_right.bold, run_right.font.name, run_right.font.size = True, 'Calibri', Pt(10.5)

            else:
                p = doc.add_paragraph()
                set_compact_spacing(p, space_after=2)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  
                add_markdown_run(p, line)

        doc.save(output_filename)
        logging.info(f"Successfully saved styled Word document as: {output_filename}")
        
    except Exception as e:
        logging.error(f"Error creating formatted resume: {e}")
        raise e

# Cleanup helper
def cleanup_files(*file_paths):
    for path in file_paths:
        if os.path.exists(path):
            try:
                os.remove(path)
                logging.info(f"Cleaned up file: {path}")
            except Exception as e:
                logging.error(f"Error cleaning up file {path}: {e}")

# ==========================================
# FASTAPI ENDPOINTS
# ==========================================

@app.post("/api/tailor-resume")
async def api_tailor_resume(
    background_tasks: BackgroundTasks,
    jd_url: str = Form(...),
    resume: UploadFile = File(...)
):
    logging.info(f"--- NEW REQUEST RECEIVED ---")
    logging.info(f"JD URL: {jd_url}")
    logging.info(f"Uploaded File: {resume.filename}")

    # NEW: Safely get filename and validate it supports PDF or TXT
    filename = resume.filename.lower()
    if not (filename.endswith('.pdf') or filename.endswith('.txt')):
        logging.warning(f"User uploaded an unsupported file type: {resume.filename}")
        raise HTTPException(status_code=400, detail="Only PDF and TXT files are supported.")

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    # NEW: Determine correct extension for temp file
    extension = ".txt" if filename.endswith('.txt') else ".pdf"
    temp_file_path = f"temp_resume_{timestamp}{extension}"
    output_docx_path = f"Tailored_Resume_{timestamp}.docx"

    # Save the uploaded file
    with open(temp_file_path, "wb") as buffer:
        shutil.copyfileobj(resume.file, buffer)

    try:
        if DEV_MODE:
            logging.info("🛠️ DEV MODE: Simulating API processing time...")
            await asyncio.sleep(3) # Simulates the delay so you can see your UI loader
            logging.info("🛠️ DEV MODE: Using mock tailored resume data.")
            tailored_text = """Dhanush Karthik
Email: dev@example.com | Phone: +91-0000000000 | LinkedIn: linkedin.com/in/dev

SECTION: PROFESSIONAL SUMMARY
Software Engineer - Backend with solid experience in developing scalable architectures. Highly skilled in **API development**, **microservices**, and **Spring Boot**.

SECTION: EXPERIENCE
Backend Developer || Jan 2024 - Present
- Designed and developed high-performance RESTful APIs using **Java** and **Spring Boot**.
- Migrated legacy monolith systems to **microservices** architecture, reducing latency by 40%.
- Integrated **PostgreSQL** and **Redis** for efficient data storage and caching.

SECTION: TECHNICAL SKILLS
**Languages:** Java, Python, JavaScript
**Frameworks/Tools:** Spring Boot, Node.js, Docker, Hibernate

SECTION: EDUCATION
Developer University || 2020 - 2024
Bachelor of Technology in Computer Science"""

        else:
            # NEW: Route text extraction based on file type
            if filename.endswith('.pdf'):
                resume_text = extract_text_from_pdf(temp_file_path)
            else:
                # Direct read for manual entry text file
                with open(temp_file_path, 'r', encoding='utf-8') as f:
                    resume_text = f.read()

            if not resume_text:
                raise Exception(f"Failed to extract text from {extension.upper()} file.")

            jd_text = await fetch_jd_from_url(jd_url)
            if not jd_text:
                raise Exception("Failed to fetch Job Description from URL.")

            tailored_text = tailor_resume_with_ai(resume_text, jd_text)
            if not tailored_text:
                raise Exception("Failed to generate tailored resume from AI.")

        # Create Word Document (This runs in both Dev and Prod modes)
        create_styled_docx(tailored_text, output_docx_path)

    except Exception as e:
        logging.error(f"Endpoint Exception Caught: {e}")
        cleanup_files(temp_file_path)
        raise HTTPException(status_code=500, detail=str(e))

    # Clean up the initial upload file
    cleanup_files(temp_file_path)
    # Schedule the DOCX to be deleted after it is sent to the user
    background_tasks.add_task(cleanup_files, output_docx_path)

    logging.info(f"--- REQUEST COMPLETED. RETURNING DOCX ---")
    return FileResponse(
        path=output_docx_path, 
        filename="Tailored_Resume.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@app.get("/")
def read_root():
    return {"message": "Resume Tailor API is running!", "dev_mode": DEV_MODE}


if __name__ == "__main__":
    import uvicorn
    # Render and other platforms provide a $PORT environment variable
    # We convert it to an integer, or default to 8000 for local dev
    port = int(os.environ.get("PORT", 8000))
    
    # Run the app on 0.0.0.0 so it is accessible externally
    uvicorn.run(app, host="0.0.0.0", port=port)